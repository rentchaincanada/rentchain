import csv
import sys
from pathlib import Path

try:
    from docx import Document
except Exception:
    Document = None  # type: ignore

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors as rl_colors
    from reportlab.pdfgen import canvas
except Exception:
    canvas = None  # type: ignore
    letter = None  # type: ignore
    rl_colors = None  # type: ignore


ROOT = Path(__file__).resolve().parents[1]
TEMPLATES_DIR = ROOT / "rentchain-frontend" / "public" / "templates"

PDF_VERSION = "Version v1.0"
MARGIN_X = 40
MARGIN_TOP = 770
PAGE_WIDTH = letter[0] if letter else 612


def ensure_deps():
    if Document is None or canvas is None:
        print("Missing dependencies. Install: python-docx reportlab", file=sys.stderr)
        sys.exit(1)


def write_docx(path: Path, title: str, paragraphs: list[str], table: list[list[str]] | None = None):
    doc = Document()
    doc.add_heading(title, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    if table:
        t = doc.add_table(rows=1, cols=len(table[0]))
        hdr = t.rows[0].cells
        for idx, cell in enumerate(hdr):
            cell.text = table[0][idx]
        for row in table[1:]:
            row_cells = t.add_row().cells
            for idx, cell in enumerate(row_cells):
                cell.text = row[idx]
    doc.save(path)


def draw_pdf_header(c, title: str):
    c.setFont("Helvetica-Bold", 12)
    c.setFillColor(rl_colors.HexColor("#111111"))
    c.drawString(MARGIN_X, MARGIN_TOP, "RENTCHAIN")
    c.setFont("Helvetica-Bold", 18)
    c.drawString(MARGIN_X, MARGIN_TOP - 22, title)
    c.setFont("Helvetica", 10)
    c.setFillColor(rl_colors.HexColor("#6b7280"))
    c.drawString(MARGIN_X, MARGIN_TOP - 40, PDF_VERSION)
    c.setFillColor(rl_colors.HexColor("#111111"))


def draw_section_title(c, title: str, y: float) -> float:
    c.setFont("Helvetica-Bold", 12)
    c.setFillColor(rl_colors.HexColor("#111111"))
    c.drawString(MARGIN_X, y, title)
    return y - 14


def draw_kv_table(c, rows: list[tuple[str, str]], y: float) -> float:
    table_width = PAGE_WIDTH - (MARGIN_X * 2)
    label_width = int(table_width * 0.38)
    value_width = table_width - label_width
    row_height = 20
    border_color = rl_colors.HexColor("#e5e7eb")
    header_fill = rl_colors.HexColor("#f3f4f6")
    text_color = rl_colors.HexColor("#111111")
    muted = rl_colors.HexColor("#6b7280")

    for label, value in rows:
        if y - row_height < 60:
            c.showPage()
            draw_pdf_header(c, c._doc.info.title or "")
            y = MARGIN_TOP - 70
        # label cell
        c.setFillColor(header_fill)
        c.rect(MARGIN_X, y - row_height, label_width, row_height, fill=1, stroke=1)
        # value cell
        c.setFillColor(rl_colors.white)
        c.rect(MARGIN_X + label_width, y - row_height, value_width, row_height, fill=1, stroke=1)
        c.setFillColor(text_color)
        c.setFont("Helvetica", 9)
        c.drawString(MARGIN_X + 8, y - 14, label)
        c.setFillColor(muted)
        c.drawString(MARGIN_X + label_width + 8, y - 14, value)
        y -= row_height
    c.setFillColor(text_color)
    return y - 10


def write_pdf(path: Path, title: str, sections: list[tuple[str, list[tuple[str, str]]]]):
    c = canvas.Canvas(str(path), pagesize=letter)
    c.setTitle(title)
    draw_pdf_header(c, title)
    y = MARGIN_TOP - 70
    for section_title, rows in sections:
        y = draw_section_title(c, section_title, y)
        y = draw_kv_table(c, rows, y)
    c.save()


def write_csv(path: Path, headers: list[str]):
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(headers)


def main():
    ensure_deps()
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

    # Notice of Entry
    write_docx(
        TEMPLATES_DIR / "Notice_of_Entry_Template.docx",
        "Notice of Entry",
        [
            "Tenant Name: {{TENANT_NAME}}",
            "Property Address: {{PROPERTY_ADDRESS}}",
            "Unit: {{UNIT_NUMBER}}",
            "Date of Notice: {{NOTICE_DATE}}",
            "Planned Entry Date/Time: {{ENTRY_DATE_TIME}}",
            "Reason for Entry: {{REASON_FOR_ENTRY}}",
            "Landlord/Manager: {{LANDLORD_NAME}}",
        ],
    )
    write_pdf(
        TEMPLATES_DIR / "Notice_of_Entry_Template.pdf",
        "NOTICE OF ENTRY",
        [
            (
                "Summary",
                [
                    ("Tenant Name", "____________________________"),
                    ("Property Address", "____________________________"),
                    ("Unit", "____________"),
                    ("Date of Notice", "____________________________"),
                    ("Planned Entry Date/Time", "____________________________"),
                    ("Reason for Entry", "____________________________"),
                    ("Landlord/Manager", "____________________________"),
                ],
            )
        ],
    )

    # Move In/Out Inspection Checklist
    write_docx(
        TEMPLATES_DIR / "Move_In_Out_Inspection_Checklist_Template.docx",
        "Move-In/Move-Out Inspection Checklist",
        [
            "Tenant Name: {{TENANT_NAME}}",
            "Property Address: {{PROPERTY_ADDRESS}}",
            "Unit: {{UNIT_NUMBER}}",
            "Inspection Type: {{MOVE_IN_OR_OUT}}",
            "Inspection Date: {{INSPECTION_DATE}}",
        ],
        table=[
            ["Area", "Condition", "Notes"],
            ["Entry / Hallway", "{{CONDITION}}", "{{NOTES}}"],
            ["Living Room", "{{CONDITION}}", "{{NOTES}}"],
            ["Kitchen", "{{CONDITION}}", "{{NOTES}}"],
            ["Bathroom", "{{CONDITION}}", "{{NOTES}}"],
            ["Bedroom", "{{CONDITION}}", "{{NOTES}}"],
        ],
    )
    write_pdf(
        TEMPLATES_DIR / "Move_In_Out_Inspection_Checklist_Template.pdf",
        "MOVE-IN / MOVE-OUT INSPECTION",
        [
            (
                "Summary",
                [
                    ("Tenant Name", "____________________________"),
                    ("Property Address", "____________________________"),
                    ("Unit", "____________"),
                    ("Inspection Type", "____________________________"),
                    ("Inspection Date", "____________________________"),
                ],
            ),
            (
                "Inspection Areas",
                [
                    ("Entry / Hallway", "____________________________"),
                    ("Living Room", "____________________________"),
                    ("Kitchen", "____________________________"),
                    ("Bathroom", "____________________________"),
                    ("Bedroom", "____________________________"),
                ],
            ),
        ],
    )

    # Rent Ledger Summary
    write_csv(
        TEMPLATES_DIR / "Rent_Ledger_Summary_Template.csv",
        ["Date", "Tenant", "Unit", "Charge Type", "Amount", "Balance"],
    )
    write_docx(
        TEMPLATES_DIR / "Rent_Ledger_Summary_Template.docx",
        "Rent Ledger Summary",
        [
            "Property: {{PROPERTY_NAME}}",
            "Period: {{PERIOD_START}} to {{PERIOD_END}}",
        ],
        table=[
            ["Date", "Tenant", "Unit", "Charge Type", "Amount", "Balance"],
            ["{{DATE}}", "{{TENANT}}", "{{UNIT}}", "{{CHARGE_TYPE}}", "{{AMOUNT}}", "{{BALANCE}}"],
        ],
    )
    write_pdf(
        TEMPLATES_DIR / "Rent_Ledger_Summary_Template.pdf",
        "RENT LEDGER SUMMARY",
        [
            (
                "Summary",
                [
                    ("Property", "____________________________"),
                    ("Period", "____________________________"),
                ],
            ),
            (
                "Ledger Columns",
                [
                    ("Date", "YYYY-MM-DD"),
                    ("Tenant", "Full name"),
                    ("Unit", "Unit number"),
                    ("Charge Type", "Rent / Fee / Credit"),
                    ("Amount", "$0.00"),
                    ("Balance", "$0.00"),
                ],
            ),
        ],
    )

    # Dispute Documentation Guide
    write_docx(
        TEMPLATES_DIR / "Dispute_Documentation_Guide_Template.docx",
        "Dispute Documentation Guide",
        [
            "Tenant: {{TENANT_NAME}}",
            "Property: {{PROPERTY_ADDRESS}}",
            "Issue Summary: {{ISSUE_SUMMARY}}",
            "Timeline:",
            "- {{DATE}}: {{EVENT}}",
            "- {{DATE}}: {{EVENT}}",
            "Supporting Evidence:",
            "- {{EVIDENCE_ITEM}}",
        ],
    )
    write_pdf(
        TEMPLATES_DIR / "Dispute_Documentation_Guide_Template.pdf",
        "DISPUTE DOCUMENTATION GUIDE",
        [
            (
                "Summary",
                [
                    ("Tenant", "____________________________"),
                    ("Property", "____________________________"),
                    ("Issue Summary", "____________________________"),
                ],
            ),
            (
                "Supporting Evidence",
                [
                    ("Timeline", "Add key dates and actions."),
                    ("Evidence Items", "List photos, emails, invoices."),
                ],
            ),
        ],
    )

    # Rental Application Checklist (Tenant)
    write_docx(
        TEMPLATES_DIR / "Rental_Application_Checklist_Tenant.docx",
        "Rental Application Checklist (Tenant)",
        [
            "Applicant: {{APPLICANT_NAME}}",
            "Email: {{APPLICANT_EMAIL}}",
            "",
            "Checklist:",
            "- Government ID",
            "- Proof of income",
            "- References",
            "- Consent to screening",
        ],
    )
    write_pdf(
        TEMPLATES_DIR / "Rental_Application_Checklist_Tenant.pdf",
        "RENTAL APPLICATION CHECKLIST",
        [
            (
                "Applicant",
                [
                    ("Name", "____________________________"),
                    ("Email", "____________________________"),
                ],
            ),
            (
                "Checklist",
                [
                    ("Government ID", "Provided / Pending"),
                    ("Proof of income", "Provided / Pending"),
                    ("References", "Provided / Pending"),
                    ("Screening consent", "Provided / Pending"),
                ],
            ),
        ],
    )

    # Late Rent Notice
    write_docx(
        TEMPLATES_DIR / "Late_Rent_Notice_Template.docx",
        "Late Rent Notice",
        [
            "Tenant Name: {{TENANT_NAME}}",
            "Property Address: {{PROPERTY_ADDRESS}}",
            "Unit: {{UNIT_NUMBER}}",
            "Rent Period: {{RENT_PERIOD}}",
            "Rent Due Date: {{RENT_DUE_DATE}}",
            "Total Rent Due: {{TOTAL_RENT_DUE}}",
            "Late Fee: {{LATE_FEE}}",
            "Other Charges: {{OTHER_CHARGES}}",
            "Total Outstanding: {{TOTAL_OUTSTANDING}}",
            "Payment Deadline: {{PAYMENT_DEADLINE}}",
            "Payment Methods: {{PAYMENT_METHODS}}",
            "Landlord/Manager: {{LANDLORD_NAME}}",
        ],
    )
    write_pdf(
        TEMPLATES_DIR / "Late_Rent_Notice_Template.pdf",
        "LATE RENT NOTICE",
        [
            (
                "Summary",
                [
                    ("Property Address", "____________________________"),
                    ("Unit", "____________________________"),
                    ("Tenant(s)", "____________________________"),
                    ("Landlord/Manager", "____________________________"),
                ],
            ),
            (
                "Amount Due",
                [
                    ("Rent Period", "____________________________"),
                    ("Rent Due Date", "____________________________"),
                    ("Total Rent Due", "____________________________"),
                    ("Late Fee", "____________________________"),
                    ("Other Charges", "____________________________"),
                    ("Total Outstanding", "____________________________"),
                ],
            ),
            (
                "Payment Instructions",
                [
                    ("Deadline", "____________________________"),
                    ("Methods", "____________________________"),
                    ("Details", "____________________________"),
                ],
            ),
            (
                "Disclaimer",
                [
                    ("Note", "This notice is provided for informational purposes only."),
                    ("Action", "Please contact your landlord/manager to resolve any disputes."),
                ],
            ),
        ],
    )

    # Tenant Notice Templates (PDF only)
    write_pdf(
        TEMPLATES_DIR / "Tenant_Notice_Templates.pdf",
        "TENANT NOTICE TEMPLATES",
        [
            (
                "Included Notices",
                [
                    ("Notice of Entry", "DOCX + PDF"),
                    ("Late Rent Notice", "DOCX + PDF"),
                    ("Lease Violation Notice", "DOCX + PDF"),
                ],
            ),
            (
                "How to Use",
                [
                    ("Edit", "Use the DOCX to customize."),
                    ("Distribute", "Export to PDF for delivery."),
                ],
            ),
        ],
    )

    # Tenant Rights Overview (PDF only)
    write_pdf(
        TEMPLATES_DIR / "Tenant_Rights_Overview.pdf",
        "TENANT RIGHTS OVERVIEW",
        [
            (
                "Overview",
                [
                    ("Purpose", "High-level summary of tenant rights."),
                    ("Note", "Refer to local jurisdiction for latest rules."),
                ],
            )
        ],
    )

    # Lease Event Log (PDF only fix)
    write_pdf(
        TEMPLATES_DIR / "Lease_Event_Log_Template.pdf",
        "LEASE EVENT LOG",
        [
            (
                "Summary",
                [
                    ("Property", "____________________________"),
                    ("Unit", "____________________________"),
                    ("Tenant", "____________________________"),
                ],
            ),
            (
                "Event Log",
                [
                    ("Date", "YYYY-MM-DD"),
                    ("Event", "Payment / Notice / Maintenance"),
                    ("Notes", "Details"),
                ],
            ),
        ],
    )

    print("Templates generated in:", TEMPLATES_DIR)


if __name__ == "__main__":
    main()
